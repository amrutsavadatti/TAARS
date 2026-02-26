"""
Text parser for plain text, markdown, and DOCX files.

Strategy per format:
- .txt / .md  : Read as UTF-8. Split into logical sections by double newline
                or markdown headings (##, ###) so chunker gets meaningful units.
- .docx       : Extract paragraph text via python-docx. Group paragraphs into
                sections using heading styles so the chunker sees the document
                structure rather than individual lines.
"""

from pathlib import Path

from backend.ingestion.models import ParsedDocument


# ---------------------------------------------------------------------------
# Plain text / Markdown
# ---------------------------------------------------------------------------

def _split_markdown_sections(text: str) -> list[str]:
    """
    Split a markdown (or plain text) document into sections.
    Splits on lines starting with # headings first; falls back to
    splitting on double blank lines.
    """
    import re

    # Try heading-based split (##+ headings)
    heading_pattern = re.compile(r"^#{1,6}\s+.+", re.MULTILINE)
    heading_positions = [m.start() for m in heading_pattern.finditer(text)]

    if len(heading_positions) >= 2:
        sections: list[str] = []
        for i, pos in enumerate(heading_positions):
            end = heading_positions[i + 1] if i + 1 < len(heading_positions) else len(text)
            section = text[pos:end].strip()
            if section:
                sections.append(section)
        return sections

    # Fallback: split on double blank lines
    sections = [s.strip() for s in re.split(r"\n{2,}", text) if s.strip()]
    return sections if sections else [text.strip()]


def parse_text(file_path: str | Path, owner_id: str) -> ParsedDocument:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    raw = file_path.read_text(encoding="utf-8", errors="replace")
    source_type = "markdown" if suffix == ".md" else "text"

    sections = _split_markdown_sections(raw)
    full_text = "\n\n".join(sections)

    return ParsedDocument(
        pages=sections,
        full_text=full_text,
        source_file=file_path.name,
        source_type=source_type,
        owner_id=owner_id,
        page_count=len(sections),
    )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _extract_docx_sections(file_path: Path) -> list[str]:
    """
    Group DOCX paragraphs into sections by heading style.
    Heading 1 / Heading 2 paragraphs start a new section.
    Non-heading paragraphs are accumulated into the current section.
    """
    import docx

    doc = docx.Document(str(file_path))
    sections: list[str] = []
    current: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        is_heading = para.style.name.startswith("Heading")

        if is_heading and current:
            sections.append("\n".join(current))
            current = [text]
        else:
            current.append(text)

    if current:
        sections.append("\n".join(current))

    return sections if sections else ["[Empty document]"]


def parse_docx(file_path: str | Path, owner_id: str) -> ParsedDocument:
    file_path = Path(file_path)
    sections = _extract_docx_sections(file_path)
    full_text = "\n\n".join(sections)

    return ParsedDocument(
        pages=sections,
        full_text=full_text,
        source_file=file_path.name,
        source_type="docx",
        owner_id=owner_id,
        page_count=len(sections),
    )
