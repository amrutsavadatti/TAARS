from __future__ import annotations

import importlib
import json
import sys

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient


def _fresh_backend_modules() -> None:
    for name in list(sys.modules):
        if name == "backend" or name.startswith("backend."):
            del sys.modules[name]


@pytest.mark.asyncio
async def test_docx_resume_import_returns_deduplicated_profile_suggestions(tmp_path, monkeypatch):
    monkeypatch.setenv("DATABASE_URL", f"sqlite+aiosqlite:///{tmp_path / 'profiles.db'}")
    _fresh_backend_modules()

    database = importlib.import_module("backend.database")
    await database.create_tables()
    import_service = importlib.import_module("backend.profile_import_service")

    model_profile = {
        "owner_name": "Test Owner",
        "experiences": [
            {
                "organization": "Acme",
                "role": "Engineer",
                "start_month": 1,
                "start_year": 2022,
                "end_month": 6,
                "end_year": 2024,
                "summary": "Built backend services.",
                "outcome": "Delivered backend services used by the product team.",
            },
            {
                "organization": "Acme",
                "role": "Engineer",
                "start_month": 1,
                "start_year": 2022,
                "end_month": 6,
                "end_year": 2024,
                "summary": "Duplicate record.",
                "outcome": "Duplicate outcome.",
            },
        ],
        "projects": [
            {
                "name": "Operations Console",
                "summary": "Built the internal console. Added audit logging and role-based access controls.",
                "problem": "Operations needed a secure internal console.",
                "contribution": "Built the console, audit logging, and access controls.",
                "outcome": "Delivered a secure operations workflow.",
                "start_month": 2,
                "start_year": 2023,
                "end_month": 8,
                "end_year": 2023,
            }
        ],
        "skills": [
            {
                "name": "Python",
                "category": "Engineering",
                "context": "Backend services",
                "evidence": "Used Python while building Acme backend services.",
            }
        ],
        "education": [],
        "certifications": [
            {
                "name": "Cloud Developer",
                "issuer": "Cloud Foundation",
                "issue_month": 5,
                "issue_year": 2024,
                "summary": "Cloud application credential.",
                "evidence": "Cloud Developer credential issued by Cloud Foundation.",
            }
        ],
        "achievements": [],
    }

    async def fake_completion(_system, _message, response_model, **_kwargs):
        return response_model.model_validate_json(json.dumps(model_profile))

    monkeypatch.setattr(import_service, "complete_structured_response", fake_completion)
    app = importlib.import_module("backend.main").app

    resume_path = tmp_path / "resume.docx"
    document = Document()
    document.add_heading("Experience", level=1)
    document.add_paragraph("Acme - Engineer, January 2022 to June 2024")
    document.add_paragraph("Built backend services with Python.")
    document.save(resume_path)

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://testserver"
    ) as client:
        with resume_path.open("rb") as resume:
            response = await client.post(
                "/api/v1/profile/import-resume",
                files={
                    "file": (
                        resume_path.name,
                        resume,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers={"X-API-Key": "dev-key"},
            )

    assert response.status_code == 200
    body = response.json()
    assert len(body["profile"]["experiences"]) == 1
    assert body["profile"]["experiences"][0]["outcome"].startswith("Delivered")
    assert body["profile"]["projects"][0]["summary"] == (
        "Built the internal console. Added audit logging and role-based access controls."
    )
    assert body["profile"]["skills"][0]["evidence"].startswith("Used Python")
    assert body["profile"]["certifications"][0]["name"] == "Cloud Developer"
    assert "experiences.0.outcome" in body["generated_fields"]
    assert any("duplicate" in warning.lower() for warning in body["warnings"])
